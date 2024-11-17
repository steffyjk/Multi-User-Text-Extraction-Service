from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from .models import DocumentDetails
from .serializers import DocumentSerializer
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import os
import requests
from .utils import send_notification

class DocumentUploadView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        serializer = DocumentSerializer(data=request.data)
        if serializer.is_valid():
            document = serializer.save()
            return Response({"message": "Document uploaded successfully, & stored in database."}, status=201)
        return Response(serializer.errors, status=400)


class ExtractTextView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        file = request.FILES.get('file')
        url = request.data.get('url')

        if not file and not url:
            return Response({"error": "Please provide a file or URL"}, status=400)

        try:
            # Handle file from URL
            if url:
                response = requests.get(url)
                file_path = f'temp/{os.path.basename(url)}'
                os.makedirs('temp', exist_ok=True)
                with open(file_path, 'wb') as f:
                    f.write(response.content)
            else:
                # Handle uploaded file
                file_path = f'temp/{file.name}'
                os.makedirs('temp', exist_ok=True)
                with open(file_path, 'wb') as f:
                    for chunk in file.chunks():
                        f.write(chunk)

            # Extract text
            if file_path.endswith('.pdf'):
                reader = PdfReader(file_path)
                text = ''.join(page.extract_text() for page in reader.pages)
            elif file_path.endswith('.docx'):
                doc = DocxDocument(file_path)
                text = '\n'.join([p.text for p in doc.paragraphs])
            else:
                return Response({"error": "Unsupported file format"}, status=400)

            # Clean up
            os.remove(file_path)

            return Response({"extracted_text": text}, status=200)

        except Exception as e:
            return Response({"error": f"Failed to extract text: {str(e)}"}, status=500)
        


class ExtractAndStoreTextView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):

        # take three types of data from request
        # 1. file 2.url 3. email
        file = request.FILES.get('file')
        url = request.data.get('url')
        email = request.data.get('email')


        # either file/url mandatory
        if not file and not url:
            return Response({"error": "Please provide a file or URL"}, status=400)
        # email is important
        if not email:
            return Response({"error": "Email is required"}, status=400)

        try:
            # Handle file from URL
            if url:

                # get the data from url
                response = requests.get(url)
                if response.status_code != 200:
                    return Response({"error": "Failed to download file from URL"}, status=400)
                
                # temp. storing that file in temp folder 
                file_path = f'temp/{os.path.basename(url)}'
                os.makedirs('temp', exist_ok=True)
                with open(file_path, 'wb') as f:
                    f.write(response.content)
            # if not url then handle the uploded file [pdf/docx]
            else:
                # Handle uploaded file
                file_path = f'temp/{file.name}'
                os.makedirs('temp', exist_ok=True)
                with open(file_path, 'wb') as f:
                    for chunk in file.chunks():
                        f.write(chunk)

            # Extract textfrom the file which we have stored in temp file path=[file_path]

            # check the file type by extention .pdf / .docx
            if file_path.endswith('.pdf'):

                # if pdf then extract text by PdfReader
                reader = PdfReader(file_path)
                extracted_text = ''.join(page.extract_text() for page in reader.pages)
            elif file_path.endswith('.docx'):

                # if docx then extract text by DocxDocument
                doc = DocxDocument(file_path)
                extracted_text = '\n'.join([p.text for p in doc.paragraphs])
            else:

                # if not .pdf or not .docx then throw error
                return Response({"error": "Unsupported file format"}, status=400)

            # Store in the database DocumentDetails table
            document = DocumentDetails.objects.create(
                file=file if file else None,
                url=url if url else None,
                extracted_text=extracted_text,
                email=email,
                status="COMPLETED",
            )

            # Cleanup / after extraction remove the file which we have stored temporiry for the extraction purpose
            os.remove(file_path)

            # Send notification on provided email
            send_notification(document.email, {
                    "message": "Document processed and stored successfully.",
                    "document_id": document.id,
                    "extracted_text": extracted_text, 
                })

            # Return success response for the response of the API
            return Response(
                {
                    "message": "Document processed and stored successfully.",
                    "document_id": document.id,
                    "extracted_text": extracted_text,  
                },
                status=201,
            )

        except Exception as e:
            return Response({"error": f"Failed to process the document: {str(e)}"}, status=500)